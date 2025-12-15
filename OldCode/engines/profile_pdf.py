import time

from gpts.gpt_agent import profileAgent

import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------
# DOC CREATION FUNC


def markdown_table_to_docx(markdown_text: str, output_path: str = None, logo_path: str = None):
    """
    Convert markdown to Docx with a logo positioned at top-left.
    - Left: -3cm indent
    - Height: Moved up by reducing header distance to 0.5cm

    If output_path is provided, saves to disk.
    Returns a BytesIO object containing the document.
    """

    lines = markdown_text.strip().split('\n')
    doc = Document()
    
    # --- LOGO POSITIONING START ---
    if logo_path:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        
        # 1. VERTICAL POSITION (Height)
        # "Header distance" is the gap from the top edge of the paper to the start of the header.
        # Default is usually ~1.27cm. Setting it to 0.5cm moves the logo UP.
        section.header_distance = Cm(0.5)
        
        # 2. HORIZONTAL POSITION (Left)
        # Align left and use negative indent to pull it into the margin.
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_para.paragraph_format.left_indent = Cm(-3)
        
        # Remove any extra spacing that might push it down
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(0)
        
        # 3. INSERT IMAGE
        run = header_para.add_run()
        try:
            # Adjust width as needed
            run.add_picture(logo_path, width=Inches(1))
        except FileNotFoundError:
            print(f"Warning: Logo file not found at {logo_path}")
    # --- LOGO POSITIONING END ---
    
    # Process text content
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Detect table start
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                line = lines[i].strip()
                if '---' not in line:
                    table_lines.append(line)
                i += 1
            
            if table_lines:
                rows = []
                for line in table_lines:
                    cells = [c.strip() for c in line.strip('|').split('|')]
                    rows.append(cells)
                
                if rows:
                    max_cols = max(len(row) for row in rows)
                    
                    # Pad rows
                    for row in rows:
                        while len(row) < max_cols:
                            row.append('')
                    
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                    doc.add_paragraph()
        
        # Headings
        elif line in [
            '1. Business Overview', '2. Key Stakeholders', '3. Revenue Split', 
            '4a. Products/Services Overview', '4b. Geographical Footprint',
            '5. Key Developments', '6. Financial Highlights', '7. Capital Structure']:
            para = doc.add_paragraph(line)
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(16)
            i += 1

        elif line in ['Summary / Interpretation', 'Sources:', 'Sources']:
            para = doc.add_paragraph(line)
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(12)
            i += 1
        
        # Bullet points
        elif line.startswith('-') or line.startswith('•'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
            i += 1
        
        # Regular paragraphs
        elif line:
            doc.add_paragraph(line)
            i += 1
        else:
            i += 1

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Optionally save to disk if output_path is provided
    if output_path:
        doc.save(output_path)
        print(f"✓ Saved to: {output_path}")

    return buffer

# ----------------------

class profile_creator():

    def __init__(self, company_name):
        self.company_name = company_name

        self.agent = profileAgent(
            company_name = self.company_name,
            k=50, 
            max_text_recall_size=35, 
            max_chars=10000,
            model='gpt-5',
        )
        self.biz_ov = ""
        self.key_stake = ""
        self.rev_split = ""
        self.prod_serv_ov = ""
        self.geo_foot = ""
        self.dev_high = ""
        self.fin_high = ""
        self.cap_stru = ""


    def _generate_section(self):

        try:
            self.biz_ov = self.agent._generate_section('GENERATE BUSINESS OVERVIEW')
        except Exception as e:
            print(f'Error at biz_ov {e}')
            pass

        try:
            self.key_stake = self.agent._generate_section('GENERATE KEY STAKEHOLDERS')
        except Exception as e:
            print(f'Error at key_stake {e}')
            pass

        try:
            self.fin_high = self.agent._generate_section('GENERATE FINANCIAL HIGHLIGHTS')
        except Exception as e:
            print(f'Error at fin_high {e}')
            pass

        try:
            self.cap_stru = self.agent._generate_section('GENERATE CAPITAL STRUCTURE')
        except Exception as e:
            print(f'Error at cap_stru {e}')
            pass

        try:
            self.rev_split = self.agent._generate_section('GENERATE REVENUE SPLIT')
        except Exception as e:
            print(f'Error at rev_split {e}')
            pass

        try:
            self.prod_serv_ov = self.agent._generate_section('GENERATE PRODUCTS SERVICES OVERVIEW')
        except Exception as e:
            print(f'Error at prod_serv_ov {e}')
            pass
        
        try:
            self.geo_foot = self.agent._generate_section('GENERATE GEO FOOTPRINT')
        except Exception as e:
            print(f'Error at geo_foot {e}')
            pass

        try:
            self.dev_high = self.agent._generate_section('GENERATE DEVELOPMENTS HIGHLIGHTS')
        except Exception as e:
            print(f'Error at dev_high {e}')
            pass

        
    def _check_sections(self):
        """
        This functions check if the variables at self. level that receives updates on _generation_section
        did indeed receive an update or still has the original value.

        IF it has original value, process that specific section again
        """
        section_mapping = {
            'biz_ov': 'GENERATE BUSINESS OVERVIEW',
            'key_stake': 'GENERATE KEY STAKEHOLDERS',
            'rev_split': 'GENERATE REVENUE SPLIT',
            'prod_serv_ov': 'GENERATE PRODUCTS SERVICES OVERVIEW',
            'geo_foot': 'GENERATE GEO FOOTPRINT',
            'dev_high': 'GENERATE DEVELOPMENTS HIGHLIGHTS',
            'fin_high': 'GENERATE FINANCIAL HIGHLIGHTS',
            'cap_stru': 'GENERATE CAPITAL STRUCTURE'
        }

        for attr_name, command in section_mapping.items():
            if getattr(self, attr_name) == "":
                try:
                    setattr(self, attr_name, self.agent._generate_section(command))
                    print(f"Regenerated section: {attr_name}")
                except Exception as e:
                    print(f'Error regenerating {attr_name}: {e}')
            

    def _unite_sections(self):
        all = "\n\n".join([self.biz_ov, self.key_stake, self.rev_split, self.prod_serv_ov, self.geo_foot, self.dev_high, self.fin_high, self.cap_stru])
        return all

# if __name__ == "__main__":

#     creator = profile_creator("Your Company Name")
#     creator._generate_section()
#     creator._check_sections()
#     all = creator._unite_sections()

#     try:
#         markdown_table_to_docx(
#             all,
#             f"{creator.company_name}.docx",
#             logo_path="logo_teneo.png"
#         )
#     except Exception as e:
#         print(e)