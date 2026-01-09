import streamlit as st
import io, re
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from OldCode.azure.blob_functions import get_file_blob
from pathlib import Path
from docx import Document
import io, tempfile, subprocess, shutil
from pathlib import Path
from typing import Optional, Tuple, Union
import os
from docx import Document
from docx.enum.text import WD_BREAK
# ---------- tile helper (the tile itself is the click target) ----------
def tile(title: str, subtitle: str, icon_svg: str, key: str):
    st.markdown(
        f"""
        <a href="?tile={key}" target="_self" style="text-decoration:none;">
          <div style="
              background:#F0F2F6; border-radius:10px; padding:14px;
              min-height:90px; display:flex; align-items:center; gap:12px;
              margin-bottom:14px; cursor:pointer;">
            <div style="width:22px;height:22px;display:flex;align-items:center;justify-content:center;">
              {icon_svg}
            </div>
            <div>
              <div style="font-weight:600; font-size:15px; color:#000;">{title}</div>
              <div style="font-size:13px; color:#555;">{subtitle}</div>
            </div>
          </div>
        </a>
        """,
        unsafe_allow_html=True
    )

def insert_finance(gpt_output, doc):
   
  # =========================
  # 1) Extract CSV + SUMMARY
  # =========================
  # Split off the summary block (everything after the blank line + heading)
  parts = gpt_output.split("\n\nSummary / Interpretation", 1)
  csv_block = parts[0].strip()

  # Ensure we start at the CSV header
  start = csv_block.find("Metric,")
  if start == -1:
      raise ValueError("CSV header 'Metric,' not found in model output.")
  csv_block = csv_block[start:]

  # Summary text (keep heading + bullets if present)
  summary_text = ""
  if len(parts) > 1:
      summary_text = "Summary / Interpretation" + parts[1].rstrip()

  # =========================
  # 2) Parse CSV to DataFrame
  # =========================
  df = pd.read_csv(io.StringIO(csv_block))
  expected_cols = {"Metric","FY24","FY23","FY22"}
  if not expected_cols.issubset(df.columns):
      raise ValueError(f"CSV columns missing. Got: {list(df.columns)}")

  # Dict: metric -> {FY24, FY23, FY22}
  csv_rows = {
      str(df.at[i, "Metric"]).strip(): {
          "FY24": df.at[i, "FY24"],
          "FY23": df.at[i, "FY23"],
          "FY22": df.at[i, "FY22"],
      }
      for i in range(len(df))
  }

  # ==============================================
  # 3) Open DOCX, find the Financial Performance table
  # ==============================================
  
  # doc_raw = get_file_blob(CONTAINER = 'templates', BLOB_NAME = 'CompaniesProfile.docx')

  # buf = io.BytesIO(doc_raw)
  # buf.seek(0) 
  

  # doc_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile (1).docx"
  # doc = Document(buf)
  doc = doc

  # ==============================================
  # 3) Open DOCX, find the Financial Performance table
  # ==============================================
  def norm(s: str) -> str:
      return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

  def find_fin_perf_table(document: Document):
      # Heuristic 1: find a table whose header row contains FY24/FY23/FY22
      for tbl in document.tables:
          if len(tbl.rows):
              header = " ".join(c.text for c in tbl.rows[0].cells)
              if all(x in norm(header) for x in ["fy24","fy23","fy22"]):
                  return tbl
      # Heuristic 2: first table after a paragraph exactly "Financial Performance"
      found_heading = False
      body = document._element.body
      for child in body.iterchildren():
          tag = child.tag.rsplit("}", 1)[-1]
          if tag == "p":
              p_text = "".join(t.text for t in child.iter() if t.tag.rsplit("}",1)[-1] == "t").strip()
              if norm(p_text) == "financialperformance":
                  found_heading = True
          elif tag == "tbl" and found_heading:
              from docx.table import Table
              return Table(child, document)
      return None

  table = find_fin_perf_table(doc)
  if table is None:
      raise RuntimeError("Could not locate the 'Financial Performance' table.")

  # ==============================================
  # 4) Map CSV metric names → rows in the template
  # ==============================================
  # --- Normalizers (use same rule for CSV and map keys) ---
  def keynorm(s: str) -> str:
      return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

  # Map CSV metric -> DOC row label (left column text in your template)
  # Use normalized keys on the left so variations in spaces/slashes/plus signs won't break it.
  metric_map_norm = {
      keynorm("Revenue (Turnover)"): "Revenue",
      keynorm("Revenue growth % (yoy)"): "Revenue Growth",
      keynorm("Gross profit"): "Gross Profit",
      keynorm("Gross margin %"): "Gross Margin",
      keynorm("EBITDA"): "EBITDA",
      keynorm("EBITDA margin %"): "EBITDA Margin",
      keynorm("Adjusted EBITDA"): "Adjusted EBITDA",

      # >>> The ones you said aren't populating <<<
      keynorm("Capex (tangible+intangible)"): "CAPEX",  # DOC label
      keynorm("Capex (tangible + intangible)"): "CAPEX",  # alt form (spaces)
      keynorm("Net Working Capital (change)"): "NET_WORK",
      keynorm("Cash Flow from Financing Activities (net)"): "Cash Flow from Financing Activities",
      keynorm("Net cash from financing activities"): "CASH_FINAN",  # alt wording
      keynorm("Total Debt (external)"): "TOTAL_DEBT",
      keynorm("Total debt (bank + lease liabilities)"): "TOTAL_DEBT",  # alt wording
      keynorm("Leverage (Net Debt/EBITDA)"): "LEVERAGE",
      keynorm("Leverage (Net Debt / EBITDA)"): "LEVERAGE",

      # Other cash flow items you have
      keynorm("Net cash from operating activities"): "Cash Flow from Operating Activities",
      keynorm("Net Working Capital (change)"): "Net Working Capital",  # duplicate on purpose (case variant)
      keynorm("Operating cash flow excl. NWC"): "Cash Flow from Operating Activities excl. Net Working Capital",
      keynorm("Other Cash Flow from Investing Activities"): "Other Cash Flow from Investing Activities",
      keynorm("Net cash from investing activities"): "Net Cash Flow from Investing Activities",
      keynorm("CFADS"): "CFADS",
      keynorm("Opening Cash"): "Opening Cash",
      keynorm("Change in Cash"): "Change in Cash",
      keynorm("Closing Cash"): "Closing Cash",
      keynorm("Bank loans outstanding"): "Total Debt",  # if your template has separate row, adjust
      keynorm("Net Debt"): "Net Debt",
  }

  # Build lookup of row labels in the DOCX (first column). You already did:
  doc_row_index = {}
  for r_idx, row in enumerate(table.rows):
      label = row.cells[0].text.strip()
      if label:
          doc_row_index[norm(label)] = r_idx  # norm = your doc normalizer (same idea as keynorm)

  # Identify FY columns in header row (handles "FY 24" vs "FY24")
  header_norm = [norm(c.text) for c in table.rows[0].cells]
  try:
      col_FY24 = header_norm.index("fy24")
      col_FY23 = header_norm.index("fy23")
      col_FY22 = header_norm.index("fy22")
  except ValueError:
      # If header is the second row in your template, try that
      header_norm = [norm(c.text) for c in table.rows[1].cells]
      col_FY24 = header_norm.index("fy24")
      col_FY23 = header_norm.index("fy23")
      col_FY22 = header_norm.index("fy22")


  # Populate table
  not_found = []

  for csv_metric, years in csv_rows.items():
      # Find the DOC row label using normalized CSV metric text
      target_label = metric_map_norm.get(keynorm(csv_metric), csv_metric)  # fall back to same text
      r_idx = doc_row_index.get(norm(target_label))  # norm() is your existing doc normalizer
      if r_idx is None:
          not_found.append(csv_metric)
          continue

      row = table.rows[r_idx]
      row.cells[col_FY24].text = str(years["FY24"])
      row.cells[col_FY23].text = str(years["FY23"])
      row.cells[col_FY22].text = str(years["FY22"])

  # ==============================================
  # 5) Populate the table from CSV
  # ==============================================
  not_found = []

  for csv_metric, years in csv_rows.items():
      target_label = metric_map_norm.get(csv_metric, csv_metric)
      r_idx = doc_row_index.get(norm(target_label))
      if r_idx is None:
          not_found.append(csv_metric)
          continue

      row = table.rows[r_idx]
      row.cells[col_FY24].text = str(years["FY24"])
      row.cells[col_FY23].text = str(years["FY23"])
      row.cells[col_FY22].text = str(years["FY22"])

  # ==============================================
  # 6) Replace the placeholder with the SUMMARY text
  #     (placeholder: [INSERT FINANCIAL PERFORMANCE SUMMARY])
  # ==============================================
  PLACEHOLDER = "[INSERT FINANCIAL PERFORMANCE SUMMARY]"

  def set_paragraph_multiline(paragraph, text: str):
      # clear runs
      for run in paragraph.runs:
          run.text = ""
      # add lines with explicit line breaks
      lines = text.splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          paragraph.add_run().add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)

  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      # search in paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # search inside tables (cells contain their own paragraphs)
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  if summary_text:
      ok = replace_placeholder(doc, PLACEHOLDER, summary_text)
      if not ok:
          print("WARNING: placeholder not found:", PLACEHOLDER)
  else:
      print("NOTE: No summary text found in GPT output (no 'Summary / Interpretation' section).")

  # ==============================================
  # 7) Save
  # ==============================================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")

  if not_found:
      print("WARNING — CSV metrics not matched to any row:")
      for m in not_found:
          print(" -", m)
  
  return doc

def insert_stakeholders(gpt_output, doc):
    
  # =========================
  # 1) Extract CSV + SUMMARY
  # =========================
  parts = gpt_output.split("\n\nSummary / Interpretation", 1)
  csv_block = parts[0].strip()

  start = csv_block.find("Metric,")
  if start == -1:
      raise ValueError("CSV header 'Metric,' not found in model output.")
  csv_block = csv_block[start:]

  summary_text = ""
  if len(parts) > 1:
      summary_text = "Summary / Interpretation" + parts[1].rstrip()

  # =========================
  # 2) Parse CSV to DataFrame
  # =========================
  df = pd.read_csv(io.StringIO(csv_block))
  expected_cols = {"Metric", "Shareholders"}
  if not expected_cols.issubset(df.columns):
      raise ValueError(f"Key Stakeholders CSV columns missing. Got: {list(df.columns)}")

  # Dict: metric -> value (right column)
  ks_rows = {
      str(df.at[i, "Metric"]).strip(): str(df.at[i, "Shareholders"]).strip()
      for i in range(len(df))
  }

  # ==============================================
  # 3) Open DOCX, find the "Key Stakeholders" table
  # ==============================================
  # doc_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile (1).docx"
  # doc_raw = get_file_blob(CONTAINER = 'templates', BLOB_NAME = 'CompaniesProfile.docx')
  # buf = io.BytesIO(doc_raw)
  # buf.seek(0) 
  # # doc_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile (1).docx"
  # doc = Document(buf)
  doc = doc

  def norm(s: str) -> str:
      return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

  def find_ks_table(document: Document):
      # Prefer a table whose header has both "Title" and "Occupants"
      for tbl in document.tables:
          if not tbl.rows:
              continue
          header = [norm(c.text) for c in tbl.rows[0].cells]
          if "title" in header and "occupants" in header:
              return tbl

      # Fallback: first table after the "Key Stakeholders" heading
      found_heading = False
      body = document._element.body
      for child in body.iterchildren():
          tag = child.tag.rsplit("}", 1)[-1]
          if tag == "p":
              p_text = "".join(t.text for t in child.iter()
                              if t.tag.rsplit("}",1)[-1] == "t").strip()
              if norm(p_text) == "keystakeholders":
                  found_heading = True
          elif tag == "tbl" and found_heading:
              from docx.table import Table
              return Table(child, document)
      return None

  ks_table = find_ks_table(doc)
  if ks_table is None:
      raise RuntimeError("Could not locate the 'Key Stakeholders' table (Title | Occupants).")

  # ==============================================
  # 4) Detect columns: label = "Title", value = "Occupants"
  # ==============================================
  def get_title_and_occupants_cols(tbl):
      # Defaults for a 2-col layout
      label_col, value_col = 0, 1

      if tbl.rows:
          header_norm = [norm(c.text) for c in tbl.rows[0].cells]
          if "title" in header_norm:
              label_col = header_norm.index("title")
          if "occupants" in header_norm:
              value_col = header_norm.index("occupants")

      # Ensure they are different; if not, force value_col to the other col
      if value_col == label_col and len(tbl.rows[0].cells) >= 2:
          value_col = 1 if label_col == 0 else 0
      return label_col, value_col

  label_col, value_col = get_title_and_occupants_cols(ks_table)

  # ==============================================
  # 5) Build row index using the Title (label) column
  # ==============================================
  row_index = {}
  for r_idx, row in enumerate(ks_table.rows):
      if not row.cells:
          continue
      # Skip header
      if r_idx == 0:
          continue
      label_text = row.cells[label_col].text.strip()
      if label_text:
          row_index[norm(label_text)] = r_idx

  # ==============================================
  # 6) Map CSV metric names → Title labels
  # ==============================================
  def keynorm(s: str) -> str:
      return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

  metric_to_title = {
      keynorm("Shareholders"): "Shareholders",
      keynorm("Management"):  "Management",
      keynorm("Lenders"):     "Lenders",
      keynorm("Auditors"):    "Auditors",
      keynorm("Advisors"):    "Advisors",
  }

  # ==============================================
  # 7) Populate the Occupants column ONLY
  # ==============================================
  not_found = []
  for metric, value in ks_rows.items():
      title_label = metric_to_title.get(keynorm(metric), metric)
      r_idx = row_index.get(norm(title_label))
      if r_idx is None:
          not_found.append(metric)
          continue
      # write into Occupants cell
      ks_table.rows[r_idx].cells[value_col].text = value

  if not_found:
      print("WARNING — missing rows for:", ", ".join(not_found))

  # ==============================================
  # 8) Replace the placeholder with the KS SUMMARY text (optional)
  #     (placeholder: [INSERT KEY STAKEHOLDERS SUMMARY])
  # ==============================================
  KS_PLACEHOLDER = "[INSERT KEY STAKEHOLDERS SUMMARY]"

  def set_paragraph_multiline(paragraph, text: str):
      for run in paragraph.runs:
          run.text = ""
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          paragraph.add_run().add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)

  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      # paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # cells
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  if summary_text:
      ok = replace_placeholder(doc, KS_PLACEHOLDER, summary_text)
      if not ok:
          print("NOTE: placeholder not found:", KS_PLACEHOLDER)

  # ==============================================
  # 9) Save
  # ==============================================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  
  print(f"Updated document written")
  return doc

def insert_capital_structure(gpt_output, doc):

  # =========================
  # 1) Extract CSV + SUMMARY  (KEEP sources)
  # =========================
  parts = gpt_output.split("\n\nSummary / Interpretation", 1)
  csv_block = parts[0].strip()

  start = csv_block.find("Metric,")
  if start == -1:
      raise ValueError("CSV header 'Metric,' not found in model output.")
  csv_block = csv_block[start:]

  summary_text = ""
  if len(parts) > 1:
      summary_text = "Summary / Interpretation" + parts[1].rstrip()

  # =========================
  # 2) Parse CSV to DataFrame
  # =========================
  df = pd.read_csv(io.StringIO(csv_block))
  expected_cols = {"Metric","FY24","FY23","FY22"}
  if not expected_cols.issubset(df.columns):
      raise ValueError(f"Capital Structure CSV columns missing. Got: {list(df.columns)}")

  csv_rows = {
      str(df.at[i, "Metric"]).strip(): {
          "FY24": str(df.at[i, "FY24"]),
          "FY23": str(df.at[i, "FY23"]),
          "FY22": str(df.at[i, "FY22"]),
      }
      for i in range(len(df))
  }

  # =========================
  # 3) Open DOCX, locate the Capital Structure table
  # =========================
  # doc_raw = get_file_blob(CONTAINER = 'templates', BLOB_NAME = 'CompaniesProfile.docx')
  # buf = io.BytesIO(doc_raw)
  # buf.seek(0) 
  # TEMPLATE_PATH = Path(__file__).parent / "templates" / "CompanyProfile (1).docx"
  # # doc_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile (1).docx"
  # doc = Document(TEMPLATE_PATH)
  doc = doc
  # ----------------- helpers -----------------
  def norm(s: str) -> str:
      return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

  def tokens(s: str) -> set:
      return set(re.findall(r"[a-z0-9]+", (s or "").lower()))

  def jaccard(a: str, b: str) -> float:
      ta, tb = tokens(a), tokens(b)
      if not ta or not tb:
          return 0.0
      inter = len(ta & tb)
      union = len(ta | tb)
      return inter / union if union else 0.0

  def find_cap_struct_table(document: Document):
      # 1) after 'Capital Structure' heading
      found_heading = False
      body = document._element.body
      for child in body.iterchildren():
          tag = child.tag.rsplit("}", 1)[-1]
          if tag == "p":
              p_text = "".join(t.text for t in child.iter()
                              if t.tag.rsplit("}",1)[-1] == "t").strip()
              if norm(p_text) == "capitalstructure":
                  found_heading = True
          elif tag == "tbl" and found_heading:
              from docx.table import Table
              return Table(child, document)

      # 2) heuristic by content
      for tbl in document.tables:
          row_texts = [" ".join(c.text for c in r.cells) for r in tbl.rows]
          joined = " ".join(row_texts)
          if all(x in norm(joined) for x in ["ebitda","leverage"]):
              return tbl
      return None

  table = find_cap_struct_table(doc)
  if table is None:
      raise RuntimeError("Could not locate the 'Capital Structure' table.")

  # Identify FY columns
  def find_fy_cols(tbl):
      for r_i in range(min(2, len(tbl.rows))):
          labels = [norm(c.text) for c in tbl.rows[r_i].cells]
          loc = {}
          for idx, txt in enumerate(labels):
              if txt == "fy24": loc["FY24"] = idx
              if txt == "fy23": loc["FY23"] = idx
              if txt == "fy22": loc["FY22"] = idx
          if {"FY24","FY23","FY22"}.issubset(loc.keys()):
              return loc["FY24"], loc["FY23"], loc["FY22"]
      if len(tbl.rows[0].cells) >= 4:
          return 1, 2, 3
      raise RuntimeError("Could not determine FY columns in Capital Structure table.")

  col_FY24, col_FY23, col_FY22 = find_fy_cols(table)

  # Build row index from first column (labels)
  doc_row_index = {}
  doc_row_labels = {}  # norm_label -> raw label (for debug)
  for r_idx, row in enumerate(table.rows):
      if not row.cells:
          continue
      label_raw = row.cells[0].text.strip()
      if label_raw:
          key = norm(label_raw)
          doc_row_index[key] = r_idx
          doc_row_labels[key] = label_raw

  # =========================
  # 4) Mapping (CSV -> DOC label), with synonyms & typo tolerance
  # =========================
  def keynorm(s: str) -> str:
      return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

  # Add broad synonyms, incl. likely template wordings
  metric_to_doc_syns = {
      keynorm("Facility Name"): [
          "Facility Name", "Name of the Facility", "Facility", "Facilities", "Facility Names",
          "Name of Facility"
      ],
      keynorm("Interest Rate"): ["Interest Rate", "Interst Rate", "Rate", "Interest"],
      keynorm("Interst Rate"):  ["Interest Rate", "Interst Rate", "Rate", "Interest"],
      keynorm("Maturity"): ["Maturity", "Final Maturity", "Maturities", "Maturity Date"],
      keynorm("Adjusted EBITDA"): ["EBITDA", "Adjusted EBITDA"],
      keynorm("Cash (Closing Cash)"): [
          "Cash (Closing Cash)", "Cash (Closing cash)", "Closing Cash",
          "Cash", "Cash and cash equivalents", "Cash & cash equivalents"
      ],
      keynorm("Net Debt"): ["Net External Debt", "Net Debt"],
      keynorm("Liquidity"): ["Liquidity"],
      keynorm("Leverage (Net Debt/EBITDA)"): ["Leverage"],
      keynorm("Leverage (Net Debt / EBITDA)"): ["Leverage"],
      keynorm("Bank loans + RCF outstanding (excl. leases)"): ["Gross External Debt", "Total External Debt"],

      keynorm("Facility B1 outstanding (GBP)"): ["Amount Outstanding"],
      keynorm("Facility B2 outstanding (GBP)"): ["Amount Outstanding"],
      keynorm("RCF drawn"): ["Amount Outstanding"],

      keynorm("RCF facility size"): [None],
      keynorm("Delayed Drawdown Facility size"): [None],
      keynorm("Bank loans due after >5 years"): [None],
      keynorm("Bank loans due within 1 year"): [None],
  }

  def smart_lookup_row_index(label_candidates):
      """
      Resolve to the best row index by:
        1) Exact normalized match
        2) Contains match (both directions)
        3) Token Jaccard similarity >= 0.5
      Returns row_index or None.
      """
      cand_norms = [norm(c) for c in label_candidates if c]

      # 1) exact
      for cn in cand_norms:
          if cn in doc_row_index:
              return doc_row_index[cn]

      # 2) contains (prefer the longest doc label match)
      best_idx = None
      best_len = -1
      for cn in cand_norms:
          for dl_norm, idx in doc_row_index.items():
              if cn and (cn in dl_norm or dl_norm in cn):
                  if len(dl_norm) > best_len:
                      best_idx, best_len = idx, len(dl_norm)
      if best_idx is not None:
          return best_idx

      # 3) token overlap
      best_idx = None
      best_score = 0.0
      for cn in cand_norms:
          for dl_norm, idx in doc_row_index.items():
              score = jaccard(cn, dl_norm)
              if score >= 0.5 and score > best_score:
                  best_idx, best_score = idx, score
      return best_idx

  # =========================
  # 5) Populate the table (incl. Facility Name / Interest Rate / Maturity)
  #    and aggregate facility amounts into 'Amount Outstanding'
  # =========================
  agg_amount = {"FY24": [], "FY23": [], "FY22": []}

  def maybe_append(prefix, v):
      v = (v or "").strip()
      if not v or v.lower() == "n.a.":
          return None
      return f"{prefix}: {v}"

  unmapped_metrics = []

  for csv_metric, years in csv_rows.items():
      mkey = keynorm(csv_metric)
      syns = metric_to_doc_syns.get(mkey, [csv_metric])

      # aggregated targets
      if any((s and norm(s) == norm("Amount Outstanding")) for s in syns if s):
          if mkey == keynorm("Facility B1 outstanding (GBP)"):
              for fy in ("FY24", "FY23", "FY22"):
                  s = maybe_append("B1", years[fy]);  agg_amount[fy].append(s) if s else None
          elif mkey == keynorm("Facility B2 outstanding (GBP)"):
              for fy in ("FY24", "FY23", "FY22"):
                  s = maybe_append("B2", years[fy]);  agg_amount[fy].append(s) if s else None
          elif mkey == keynorm("RCF drawn"):
              for fy in ("FY24", "FY23", "FY22"):
                  s = maybe_append("RCF drawn", years[fy]);  agg_amount[fy].append(s) if s else None
          continue

      r_idx = smart_lookup_row_index(syns)
      if r_idx is None:
          unmapped_metrics.append(csv_metric)
          continue

      row = table.rows[r_idx]
      row.cells[col_FY24].text = years["FY24"]
      row.cells[col_FY23].text = years["FY23"]
      row.cells[col_FY22].text = years["FY22"]

  # Aggregated 'Amount Outstanding'
  amount_row_idx = smart_lookup_row_index(["Amount Outstanding"])
  if amount_row_idx is not None:
      row = table.rows[amount_row_idx]
      row.cells[col_FY24].text = "; ".join(agg_amount["FY24"]) if agg_amount["FY24"] else "n.a."
      row.cells[col_FY23].text = "; ".join(agg_amount["FY23"]) if agg_amount["FY23"] else "n.a."
      row.cells[col_FY22].text = "; ".join(agg_amount["FY22"]) if agg_amount["FY22"] else "n.a."

  # =========================
  # 6) Insert the full SUMMARY (including Sources) — no duplicates
  # =========================
  PLACEHOLDER = "[INSERT CAPITAL STRUCTURE SUMMARY]"
  HEADING_TEXT = "Summary / Interpretation"

  def set_paragraph_multiline(paragraph, text: str):
      for run in paragraph.runs:  # clear
          run.text = ""
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          paragraph.add_run().add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)

  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      # paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # tables
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  def find_all_summary_paragraphs(document: Document, heading_text: str):
      anchors = []
      for p in document.paragraphs:
          if heading_text in p.text:
              anchors.append(p)
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if heading_text in p.text:
                          anchors.append(p)
      return anchors

  insert_done = False
  if summary_text:
      insert_done = replace_placeholder(doc, PLACEHOLDER, summary_text)

  if summary_text and not insert_done:
      anchors = find_all_summary_paragraphs(doc, HEADING_TEXT)
      if anchors:
          # overwrite first and remove extras
          set_paragraph_multiline(anchors[0], summary_text)
          for dup in anchors[1:]:
              dup._element.getparent().remove(dup._element)
      else:
          p = doc.add_paragraph()
          set_paragraph_multiline(p, summary_text)
          # ensure no accidental multiples
          anchors = find_all_summary_paragraphs(doc, HEADING_TEXT)
          for dup in anchors[1:]:
              dup._element.getparent().remove(dup._element)

  # =========================
  # 7) Save + (optional) debug
  # =========================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")

  if unmapped_metrics:
      print("NOTE — CSV metrics that couldn't be matched to any row (check your template labels):")
      for m in unmapped_metrics:
          print(" -", m)
  
  return doc

def insert_biz_overview(gpt_output):

  # =========================
  # 1) Open DOCX
  # =========================
  # doc_raw = get_file_blob(CONTAINER = 'templates', BLOB_NAME = 'CompaniesProfile.docx')
  # buf = io.BytesIO(doc_raw)
  # buf.seek(0) 
  # # doc_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile (1).docx"
  # doc = Document(buf)

  ROOT = Path(__file__).resolve().parents[2]  # …/Azure-OnePager
  TEMPLATE_PATH = ROOT / "templates" / "CompanyProfile (1).docx"
  doc = Document(TEMPLATE_PATH)
  # TEMPLATE_PATH = Path(__file__).parent / "templates" / "CompanyProfile (1).docx"
  # doc_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile (1).docx"
 

  PLACEHOLDER = "[INSERT BUSINESS OVERVIEW]"

  def set_paragraph_multiline(paragraph, text: str):
      """Replace a paragraph's text with multi-line content, preserving line breaks."""
      # clear existing runs
      for run in paragraph.runs:
          run.text = ""
      # write lines with explicit line breaks
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          r = paragraph.add_run()
          r.add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)


  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      """Find placeholder in paragraphs/cells and replace it with new_text (multiline)."""
      # plain paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # inside tables
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  # =========================
  # 2) Replace the placeholder
  # =========================
  ok = replace_placeholder(doc, PLACEHOLDER, gpt_output)
  if not ok:
      print(f"WARNING: placeholder not found: {PLACEHOLDER}")

  # =========================
  # 3) Save
  # =========================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")
  return doc

def insert_revenue_split(gpt_output, doc):

  PLACEHOLDER = "[INSERT REVENUE SPLIT]"

  def set_paragraph_multiline(paragraph, text: str):
      """Replace a paragraph's text with multi-line content, preserving line breaks."""
      # clear existing runs
      for run in paragraph.runs:
          run.text = ""
      # write lines with explicit line breaks
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          r = paragraph.add_run()
          r.add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)


  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      """Find placeholder in paragraphs/cells and replace it with new_text (multiline)."""
      # plain paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # inside tables
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  # =========================
  # 2) Replace the placeholder
  # =========================
  ok = replace_placeholder(doc, PLACEHOLDER, gpt_output)
  if not ok:
      print(f"WARNING: placeholder not found: {PLACEHOLDER}")

  # =========================
  # 3) Save
  # =========================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")
  return doc

def insert_services_overview(gpt_output, doc):

  PLACEHOLDER = "[INSERT SERVICES OVERVIEW]"

  def set_paragraph_multiline(paragraph, text: str):
      """Replace a paragraph's text with multi-line content, preserving line breaks."""
      # clear existing runs
      for run in paragraph.runs:
          run.text = ""
      # write lines with explicit line breaks
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          r = paragraph.add_run()
          r.add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)


  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      """Find placeholder in paragraphs/cells and replace it with new_text (multiline)."""
      # plain paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # inside tables
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  # =========================
  # 2) Replace the placeholder
  # =========================
  ok = replace_placeholder(doc, PLACEHOLDER, gpt_output)
  if not ok:
      print(f"WARNING: placeholder not found: {PLACEHOLDER}")

  # =========================
  # 3) Save
  # =========================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")
  return doc

def insert_geo_footprint(gpt_output, doc):
  PLACEHOLDER = "[INSERT GEO FOOTPRINT]"

  def set_paragraph_multiline(paragraph, text: str):
      """Replace a paragraph's text with multi-line content, preserving line breaks."""
      # clear existing runs
      for run in paragraph.runs:
          run.text = ""
      # write lines with explicit line breaks
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          r = paragraph.add_run()
          r.add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)


  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      """Find placeholder in paragraphs/cells and replace it with new_text (multiline)."""
      # plain paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # inside tables
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  # =========================
  # 2) Replace the placeholder
  # =========================
  ok = replace_placeholder(doc, PLACEHOLDER, gpt_output)
  if not ok:
      print(f"WARNING: placeholder not found: {PLACEHOLDER}")

  # =========================
  # 3) Save
  # =========================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")
  return doc

def insert_development_highlights(gpt_output, doc):
  PLACEHOLDER = "[INSERT KEY DEVELOPMENTS]"

  def set_paragraph_multiline(paragraph, text: str):
      """Replace a paragraph's text with multi-line content, preserving line breaks."""
      # clear existing runs
      for run in paragraph.runs:
          run.text = ""
      # write lines with explicit line breaks
      lines = (text or "").splitlines()
      if not lines:
          return
      paragraph.add_run(lines[0])
      for ln in lines[1:]:
          r = paragraph.add_run()
          r.add_break(WD_BREAK.LINE)
          paragraph.add_run(ln)


  def replace_placeholder(document: Document, placeholder: str, new_text: str) -> bool:
      """Find placeholder in paragraphs/cells and replace it with new_text (multiline)."""
      # plain paragraphs
      for p in document.paragraphs:
          if placeholder in p.text:
              set_paragraph_multiline(p, new_text)
              return True
      # inside tables
      for tbl in document.tables:
          for row in tbl.rows:
              for cell in row.cells:
                  for p in cell.paragraphs:
                      if placeholder in p.text:
                          set_paragraph_multiline(p, new_text)
                          return True
      return False

  # =========================
  # 2) Replace the placeholder
  # =========================
  ok = replace_placeholder(doc, PLACEHOLDER, gpt_output)
  if not ok:
      print(f"WARNING: placeholder not found: {PLACEHOLDER}")

  # =========================
  # 3) Save
  # =========================
  # out_path = "/Users/felipesilverio/Documents/GitHub/Azure-OnePager/CompanyProfile2.docx"
  # doc.save(out_path)
  print(f"Updated document written")
  return doc

def _docx_bytes_to_pdf_bytes_with_docx2pdf(docx_bytes: bytes) -> Optional[bytes]:
    """Try converting via docx2pdf (requires MS Word on Windows/macOS). Returns None if unavailable/fails."""
    try:
        from docx2pdf import convert  # type: ignore
    except Exception:
        return None

    try:
        with tempfile.TemporaryDirectory() as td:
            in_path  = Path(td) / "doc.docx"
            out_path = Path(td) / "doc.pdf"
            in_path.write_bytes(docx_bytes)
            convert(str(in_path), str(out_path))  # uses Word/Automator
            return out_path.read_bytes()
    except Exception as e:
        print("[docx2pdf] conversion failed:", repr(e))
        return None
    
def _docx_bytes_to_pdf_bytes_with_lo(docx_bytes: bytes) -> Optional[bytes]:
    """Try converting via LibreOffice (soffice --headless). Returns None if not installed/fails."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        with tempfile.TemporaryDirectory() as td:
            in_path  = Path(td) / "doc.docx"
            out_dir  = Path(td) / "out"
            out_dir.mkdir(parents=True, exist_ok=True)
            in_path.write_bytes(docx_bytes)
            # Convert
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(in_path)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            out_path = out_dir / "doc.pdf"
            if out_path.exists():
                return out_path.read_bytes()
            return None
    except Exception as e:
        print("[LibreOffice] exception:", repr(e))
        return None
    
def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> Optional[bytes]:
    """Best-effort conversion with two strategies."""
    pdf = _docx_bytes_to_pdf_bytes_with_docx2pdf(docx_bytes)
    if pdf:
        return pdf
    
    pdf = _docx_bytes_to_pdf_bytes_with_lo(docx_bytes)
    if pdf:
        return _docx_bytes_to_pdf_bytes_with_lo(docx_bytes)
    
    return 'Unable to transform to pdf'

def save_docx_to_pdf_via_libreoffice(doc, pdf_path: str):
    tmpdir = tempfile.mkdtemp()
    docx_path = os.path.join(tmpdir, "report.docx")
    doc.save(docx_path)

    outdir = str(Path(pdf_path).parent)
    os.makedirs(outdir, exist_ok=True)

    # Convert using LibreOffice in headless mode
    cmd = [
        "soffice", "--headless", "--nologo",
        "--convert-to", "pdf:writer_pdf_Export",
        "--outdir", outdir, docx_path
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    produced = os.path.join(outdir, "report.pdf")
    if res.returncode != 0 or not os.path.exists(produced):
        raise RuntimeError(f"LibreOffice failed:\nSTDOUT:\n{res.stdout}\nSTDERR:\n{res.stderr}")

    shutil.move(produced, pdf_path)
    shutil.rmtree(tmpdir, ignore_errors=True)